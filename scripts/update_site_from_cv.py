#!/usr/bin/env python3
"""Regenerate the personal site data from Hanall Sung's Word CV."""

from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
CV_PATH = ROOT / "CV_Hanall Sung_with Index.docx"
SITE_DIR = ROOT / "docs"
JSON_PATH = SITE_DIR / "data" / "cv.json"
JS_PATH = SITE_DIR / "assets" / "cv-data.js"


def clean(text: str) -> str:
    return " ".join((text or "").split())


def unique_cells(row) -> list[str]:
    items: list[str] = []
    for cell in row.cells:
        value = clean(cell.text)
        if value and value not in items:
            items.append(value)
    return items


def is_yearish(text: str) -> bool:
    return bool(re.search(r"(19|20)\d{2}|present|under\b|accepted|in progress", text, re.I))


def split_entry(text: str) -> dict[str, str]:
    text = clean(text)
    link = ""
    doi = ""
    link_match = re.search(r"https?://\S+", text)
    if link_match:
        link = link_match.group(0).rstrip(").,")
        doi_match = re.search(r"10\.\d{4,9}/[^\s\]]+", link)
        if doi_match:
            doi = doi_match.group(0).rstrip(").,")

    year_match = re.search(r"\((accepted|under[^)]*|in progress|20\d{2}|19\d{2})\)", text, re.I)
    year = year_match.group(1) if year_match else ""

    title = text
    if year_match:
        after_year = text[year_match.end() :].strip()
        title = after_year.split(".", 1)[0].strip() or text

    return {
        "text": text,
        "title": title,
        "year": year,
        "link": link,
        "doi": doi,
        "corresponding_author": "Sung, H.*" in text or "Sung, H. *" in text,
    }


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def split_title_source(text: str) -> dict[str, str]:
    parts = re.split(r"\s+Funding source:\s+", text, maxsplit=1)
    title = clean(parts[0])
    funding_source = ""
    role = ""
    details = clean(parts[1]) if len(parts) > 1 else ""
    if details:
        role_parts = re.split(r"\s+Role:\s+", details, maxsplit=1)
        funding_source = clean(role_parts[0])
        role = clean(role_parts[1]) if len(role_parts) > 1 else ""
    return {
        "title": title,
        "funding_source": funding_source,
        "role": role,
        "text": text,
    }


def date_rows(table) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for row in table.rows:
        cells = unique_cells(row)
        if len(cells) >= 2 and cells[0] and cells[1]:
            rows.append({"date": cells[0], "text": cells[1]})
    return rows


def parse_table_one(table) -> dict[str, list]:
    data = {
        "appointments": [],
        "education": [],
        "honors": [],
    }
    section = ""
    for row in table.rows:
        cells = unique_cells(row)
        if not cells:
            continue
        label = cells[0]
        upper = label.upper()
        if upper in {"PROFESSIONAL APPOINTMENT", "EDUCATION", "AWARDS • HONORS • SCHOLARSHIP"}:
            section = upper
            continue
        if section == "PROFESSIONAL APPOINTMENT" and len(cells) == 1:
            data["appointments"].append({"title": label})
        elif section == "EDUCATION" and len(cells) == 1:
            data["education"].append({"degree": label})
        elif section == "AWARDS • HONORS • SCHOLARSHIP" and len(cells) >= 2:
            data["honors"].append({"date": cells[0], "text": cells[1]})
    return data


def parse_experience_teaching(table) -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    research: list[dict[str, str]] = []
    teaching: list[dict[str, str]] = []
    section = "research"
    institution = ""
    for row in table.rows:
        cells = unique_cells(row)
        if not cells:
            continue
        if cells[0] == "RESEARCH EXPERIENCE":
            section = "research"
            continue
        if cells[0] == "TEACHING EXPERIENCE":
            section = "teaching"
            institution = ""
            continue
        if len(cells) == 1 and section == "teaching":
            institution = cells[0]
            continue
        if len(cells) >= 2:
            item = {"date": cells[0], "text": cells[1]}
            if section == "teaching":
                item["institution"] = institution
                teaching.append(item)
            else:
                research.append(item)
    return research, teaching


def parse_advising(table) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    group = ""
    for row in table.rows:
        cells = unique_cells(row)
        if not cells or cells[0] == "ADVISING & MENTORING EXPERIENCE":
            continue
        if len(cells) == 1:
            group = cells[0]
        elif len(cells) >= 2:
            rows.append({"group": group, "date": cells[0], "text": cells[1]})
    return rows


def parse_docx() -> dict:
    doc = Document(CV_PATH)
    paragraphs = [clean(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]

    profile = {
        "name": paragraphs[0],
        "affiliation": paragraphs[1],
        "address": paragraphs[2],
        "contact": paragraphs[3],
        "email": "hanallsung@utk.edu",
        "interests": paragraphs[4],
    }

    t1 = parse_table_one(doc.tables[1])
    extramural = []
    for item in date_rows(doc.tables[3]):
        item.update(split_title_source(item["text"]))
        extramural.append(item)
    intramural = []
    for item in date_rows(doc.tables[4]):
        item.update(split_title_source(item["text"]))
        intramural.append(item)

    section = ""
    publications: dict[str, list] = {
        "published": [],
        "under_review": [],
        "in_preparation": [],
        "book_chapters": [],
        "conference_proceedings": [],
        "invited_talks": [],
    }
    skills: list[str] = []
    language = ""
    certification = ""

    heading_map = {
        "Published": "published",
        "Manuscripts submitted for publication": "under_review",
        "In preparation": "in_preparation",
        "Book Chapter": "book_chapters",
        "Papers published in refereed conference proceedings": "conference_proceedings",
    }

    for block in iter_blocks(doc):
        if isinstance(block, Table):
            table_text = clean(" ".join(cell.text for row in block.rows for cell in row.cells))
            if "INVITED TALKS AND PRESENTATIONS" in table_text:
                section = "invited_talks"
            continue

        text = clean(block.text)
        if not text:
            continue
        if text in heading_map:
            section = heading_map[text]
            continue
        if text in {
            "Hanall Sung, Ph.D.",
            profile["affiliation"],
            profile["address"],
            profile["contact"],
            profile["interests"],
            "Extramural Grants",
            "Intramural Grants",
            "Articles published in refereed journals",
            "* indicates corresponding author",
            "Peer-reviewed Journal Reviewer",
            "Annual Conference Proposal Reviewer",
            "Conference Committee Roles",
            "Scholarly Community Involvement",
            "Research Methods and Software",
            "Language",
            "Certification",
        }:
            if text in {
                "Peer-reviewed Journal Reviewer",
                "Annual Conference Proposal Reviewer",
                "Conference Committee Roles",
                "Scholarly Community Involvement",
            }:
                section = ""
            elif text == "Research Methods and Software":
                section = "skills"
            elif text == "Language":
                section = "language"
            elif text == "Certification":
                section = "certification"
            continue
        if section in publications and (is_yearish(text) or "doi" in text.lower()):
            publications[section].append(split_entry(text))
        elif section == "skills":
            skills.append(text)
        elif section == "language":
            language = text
        elif section == "certification":
            certification = text

    research_experience, teaching = parse_experience_teaching(doc.tables[7])

    data = {
        "profile": profile,
        "appointments": t1["appointments"],
        "education": t1["education"],
        "honors": t1["honors"],
        "grants": {
            "extramural": extramural,
            "intramural": intramural,
        },
        "publications": publications,
        "research_experience": research_experience,
        "teaching": teaching,
        "advising": parse_advising(doc.tables[8]),
        "affiliations": {
            "institutional": t1["appointments"],
            "scholarly_communities": date_rows(doc.tables[14]),
        },
        "service": {
            "editorial": date_rows(doc.tables[10]),
            "conference_reviewing": date_rows(doc.tables[11]),
            "professional": date_rows(doc.tables[13]),
            "community": date_rows(doc.tables[14]),
        },
        "skills": skills,
        "language": language,
        "certification": certification,
        "source": {
            "cv_file": CV_PATH.name,
            "generated_at": datetime.now().isoformat(timespec="seconds"),
        },
    }
    return data


def write_outputs(data: dict) -> None:
    JSON_PATH.parent.mkdir(parents=True, exist_ok=True)
    JS_PATH.parent.mkdir(parents=True, exist_ok=True)
    JSON_PATH.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    JS_PATH.write_text(
        "window.CV_DATA = " + json.dumps(data, indent=2, ensure_ascii=False) + ";\n",
        encoding="utf-8",
    )


def main() -> None:
    data = parse_docx()
    write_outputs(data)
    total_publications = sum(len(items) for items in data["publications"].values())
    print(f"Generated {JSON_PATH.relative_to(ROOT)}")
    print(f"Generated {JS_PATH.relative_to(ROOT)}")
    print(f"Parsed {total_publications} publication/presentation entries")


if __name__ == "__main__":
    main()
